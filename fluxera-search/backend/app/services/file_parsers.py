from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader


# Plain-text types decoded directly.
_PLAINTEXT_EXTENSIONS = {".txt", ".md", ".py", ".rst", ".json", ".csv", ".yaml", ".yml", ".toml"}
SUPPORTED_EXTENSIONS = _PLAINTEXT_EXTENSIONS | {".pdf", ".docx", ".html", ".htm"}


def parse_file_bytes(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: {suffix}. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if suffix in _PLAINTEXT_EXTENSIONS:
        return data.decode("utf-8", errors="ignore")

    if suffix == ".pdf":
        return _parse_pdf(data)

    if suffix == ".docx":
        return _parse_docx(data)

    # HTML / HTM
    return _parse_html(data)


def _parse_pdf(data: bytes) -> str:
    from io import BytesIO

    try:
        reader = PdfReader(BytesIO(data))
    except Exception as exc:
        raise ValueError(
            "Unable to open PDF. The file may be encrypted or corrupted."
        ) from exc

    pages: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        pages.append(text)

    extracted = "\n\n".join(p for p in pages if p.strip())
    if not extracted.strip():
        raise ValueError(
            "No text found in PDF. The file may be image-only (scanned). "
            "Please use a PDF with selectable text or convert it with OCR first."
        )
    return extracted


def _parse_docx(data: bytes) -> str:
    from io import BytesIO

    doc = DocxDocument(BytesIO(data))
    parts: list[str] = []

    # Paragraphs (preserves heading/body text order).
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables — extract cell text row by row.
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def _parse_html(data: bytes) -> str:
    soup = BeautifulSoup(data.decode("utf-8", errors="ignore"), "html.parser")
    # Remove script, style, and nav noise before extracting text.
    for tag in soup(["script", "style", "noscript", "nav", "footer", "header"]):
        tag.decompose()
    return soup.get_text("\n", strip=True)
